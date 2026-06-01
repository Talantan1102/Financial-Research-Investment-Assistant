"""知识库管理路由"""

import asyncio
import logging
import os
import re
import shutil
import uuid as _uuid
from uuid import UUID

from fastapi import APIRouter, BackgroundTasks, Depends, File, HTTPException, UploadFile, status
from sqlalchemy.orm import Session

from app.core.database import get_db
from app.models.knowledge import Document, KnowledgeBase
from app.models.user import User
from app.router._upload_utils import ALLOWED_EXTENSIONS, get_file_extension  # C72: shared SSOT
from app.router.auth_router import get_current_user_required
from app.schemas.knowledge import (
    DocumentResponse,
    DocumentUploadResponse,
    KnowledgeBaseCreate,
    KnowledgeBaseResponse,
    KnowledgeBaseUpdate,
    KnowledgeBaseWithDocuments,
)

router = APIRouter(prefix="/knowledge-bases", tags=["知识库管理"])

logger = logging.getLogger(__name__)

# 文件上传目录
UPLOAD_DIR = "/tmp/knowledge_uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

# 直读为文本的扩展名(其余走 pdf/docx 解析或二进制兜底)
_TEXT_EXTENSIONS = {
    "txt", "md", "markdown", "json", "csv", "py", "js", "ts",
    "jsx", "tsx", "yaml", "yml", "xml", "log", "html", "htm",
}


def _extract_text(file_path: str, filename: str) -> str:
    """从上传文件抽取纯文本。txt/code 直读;pdf 走 pdfplumber;docx 走 python-docx。"""
    ext = os.path.splitext(filename)[1].lower().lstrip(".")
    if ext in _TEXT_EXTENSIONS:
        with open(file_path, "rb") as f:
            raw = f.read()
        # utf-8 优先,gb18030 兜中文遗留编码;都失败再用 replace 出可见替换符
        # (不放 latin-1:它对任意字节都不抛错,会把中文硬解成 mojibake 污染 embedding)。
        for enc in ("utf-8", "gb18030"):
            try:
                return raw.decode(enc)
            except UnicodeDecodeError:
                continue
        return raw.decode("utf-8", errors="replace")
    if ext == "pdf":
        import pdfplumber

        parts: list[str] = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                t = page.extract_text() or ""
                if t.strip():
                    parts.append(t)
        return "\n\n".join(parts)
    if ext in ("docx", "doc"):
        import docx

        d = docx.Document(file_path)
        return "\n".join(p.text for p in d.paragraphs if p.text.strip())
    # 兜底:按文本读
    with open(file_path, "rb") as f:
        return f.read().decode("utf-8", errors="replace")


def _chunk_text(text: str, chunk_size: int = 600, overlap: int = 80) -> list[str]:
    """段落感知的简单切块:按空行切段,贪心打包到 ~chunk_size;超长段硬切。"""
    text = text.strip()
    if not text:
        return []
    paras = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    chunks: list[str] = []
    buf = ""
    for p in paras:
        if len(buf) + len(p) + 1 <= chunk_size:
            buf = f"{buf}\n{p}".strip()
        else:
            if buf:
                chunks.append(buf)
                buf = ""
            if len(p) <= chunk_size:
                buf = p
            else:
                step = max(chunk_size - overlap, 1)
                for i in range(0, len(p), step):
                    chunks.append(p[i : i + chunk_size])
    if buf:
        chunks.append(buf)
    return chunks


def kb_to_response(kb: KnowledgeBase) -> KnowledgeBaseResponse:
    """将知识库模型转换为响应"""
    return KnowledgeBaseResponse(
        id=str(kb.id),
        name=kb.name,
        description=kb.description,
        document_count=kb.document_count or 0,
        created_at=kb.created_at,
        updated_at=kb.updated_at,
    )


def doc_to_response(doc: Document) -> DocumentResponse:
    """将文档模型转换为响应"""
    return DocumentResponse(
        id=str(doc.id),
        knowledge_base_id=str(doc.knowledge_base_id),
        filename=doc.filename,
        file_type=doc.file_type,
        file_size=doc.file_size,
        status=doc.status,
        chunk_count=doc.chunk_count or 0,
        error_message=doc.error_message,
        created_at=doc.created_at,
        updated_at=doc.updated_at,
    )


async def process_document(document_id: str, file_path: str, kb_id: str, db_session_factory):
    """后台处理文档：解析 → 分块 → qwen embed → Milvus(kb_{kb_id})。

    v1.x 修复：原实现 import `app.service.docmind_service`，该模块依赖**未声明**的
    `alibabacloud_docmind_api20220711` SDK，import 即 ModuleNotFoundError，后台任务
    一进来就抛错、文档永久卡 "pending" 且不向用户暴露（违反 fail-loud）。改走本地
    pdfplumber/直读解析 + EMBEDDING_MODE(qwen) + MilvusService.insert_documents，
    写入与 get_document_chunks 一致的 `kb_{kb_id}` collection。错误写入
    doc.error_message 让前端可见。
    """
    from app.service.milvus_service import get_milvus_service
    from app.services.embedding_factory import build_embedding_service_from_env

    db = db_session_factory()
    try:
        doc = db.query(Document).filter(Document.id == document_id).first()
        if not doc:
            return

        doc.status = "processing"
        db.commit()

        try:
            # C73: KB UUID 作为 collection 名，避免同名 KB 跨用户撞 collection。
            index_name = f"kb_{kb_id}".replace("-", "_")

            # 解析是阻塞 IO(pdfplumber/文件读),放线程池避免堵 async 事件循环。
            text = await asyncio.to_thread(_extract_text, file_path, doc.filename)
            chunks = _chunk_text(text)
            if not chunks:
                raise ValueError("未能从文档提取到任何文本内容")

            embedder = build_embedding_service_from_env()
            vectors = await embedder.embed(chunks)  # 已是 async(内部 to_thread 调 dashscope)

            documents = [
                {
                    "id": str(_uuid.uuid4()),
                    "doc_id": str(doc.id),
                    "kb_id": str(kb_id),
                    "filename": doc.filename,
                    "content": chunk,
                    "chunk_index": i,
                    "vector": vector,
                }
                for i, (chunk, vector) in enumerate(zip(chunks, vectors, strict=True))
            ]
            # get_milvus_service() 首调会 connect、insert/flush 都是阻塞 pymilvus RPC,
            # 一并放线程池,避免摄取大文档时堵塞整个 FastAPI 事件循环。
            milvus = await asyncio.to_thread(get_milvus_service)
            count = await asyncio.to_thread(milvus.insert_documents, index_name, documents)

            doc.status = "completed"
            doc.chunk_count = count
            doc.error_message = None
            logger.info(
                "KB ingest done doc=%s chunks=%d collection=%s", doc.id, count, index_name
            )

        except Exception as e:  # noqa: BLE001 — 错误隔离 + fail-loud 写库
            logger.exception("KB ingest failed doc=%s", document_id)
            doc.status = "failed"
            doc.error_message = f"{type(e).__name__}: {e}"

        db.commit()

    finally:
        db.close()
        # 清理临时文件
        if os.path.exists(file_path):
            os.remove(file_path)


@router.get("", response_model=list[KnowledgeBaseResponse])
async def get_knowledge_bases(
    current_user: User = Depends(get_current_user_required),
    db: Session = Depends(get_db),
):
    """获取用户的知识库列表"""
    kbs = (
        db.query(KnowledgeBase)
        .filter(KnowledgeBase.user_id == current_user.id)
        .order_by(KnowledgeBase.updated_at.desc())
        .all()
    )

    return [kb_to_response(kb) for kb in kbs]


@router.post("", response_model=KnowledgeBaseResponse, status_code=status.HTTP_201_CREATED)
async def create_knowledge_base(
    kb_data: KnowledgeBaseCreate,
    current_user: User = Depends(get_current_user_required),
    db: Session = Depends(get_db),
):
    """创建知识库"""
    # 检查是否已存在同名知识库
    existing = (
        db.query(KnowledgeBase)
        .filter(KnowledgeBase.user_id == current_user.id, KnowledgeBase.name == kb_data.name)
        .first()
    )

    if existing:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="已存在同名知识库")

    kb = KnowledgeBase(
        user_id=current_user.id,
        name=kb_data.name,
        description=kb_data.description,
    )
    db.add(kb)
    db.commit()
    db.refresh(kb)

    return kb_to_response(kb)


@router.get("/{kb_id}", response_model=KnowledgeBaseWithDocuments)
async def get_knowledge_base(
    kb_id: str,
    current_user: User = Depends(get_current_user_required),
    db: Session = Depends(get_db),
):
    """获取知识库详情（包含文档列表）"""
    try:
        kb_uuid = UUID(kb_id)
    except ValueError:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="无效的知识库ID格式")

    kb = (
        db.query(KnowledgeBase)
        .filter(KnowledgeBase.id == kb_uuid, KnowledgeBase.user_id == current_user.id)
        .first()
    )

    if not kb:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="知识库不存在")

    documents = (
        db.query(Document)
        .filter(Document.knowledge_base_id == kb.id)
        .order_by(Document.created_at.desc())
        .all()
    )

    return KnowledgeBaseWithDocuments(
        id=str(kb.id),
        name=kb.name,
        description=kb.description,
        document_count=kb.document_count or 0,
        created_at=kb.created_at,
        updated_at=kb.updated_at,
        documents=[doc_to_response(doc) for doc in documents],
    )


@router.put("/{kb_id}", response_model=KnowledgeBaseResponse)
async def update_knowledge_base(
    kb_id: str,
    kb_data: KnowledgeBaseUpdate,
    current_user: User = Depends(get_current_user_required),
    db: Session = Depends(get_db),
):
    """更新知识库"""
    try:
        kb_uuid = UUID(kb_id)
    except ValueError:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="无效的知识库ID格式")

    kb = (
        db.query(KnowledgeBase)
        .filter(KnowledgeBase.id == kb_uuid, KnowledgeBase.user_id == current_user.id)
        .first()
    )

    if not kb:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="知识库不存在")

    if kb_data.name is not None:
        # 检查是否与其他知识库重名
        existing = (
            db.query(KnowledgeBase)
            .filter(
                KnowledgeBase.user_id == current_user.id,
                KnowledgeBase.name == kb_data.name,
                KnowledgeBase.id != kb_uuid,
            )
            .first()
        )
        if existing:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="已存在同名知识库")
        kb.name = kb_data.name

    if kb_data.description is not None:
        kb.description = kb_data.description

    db.commit()
    db.refresh(kb)

    return kb_to_response(kb)


@router.delete("/{kb_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_knowledge_base(
    kb_id: str,
    current_user: User = Depends(get_current_user_required),
    db: Session = Depends(get_db),
):
    """删除知识库"""
    try:
        kb_uuid = UUID(kb_id)
    except ValueError:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="无效的知识库ID格式")

    kb = (
        db.query(KnowledgeBase)
        .filter(KnowledgeBase.id == kb_uuid, KnowledgeBase.user_id == current_user.id)
        .first()
    )

    if not kb:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="知识库不存在")

    db.delete(kb)
    db.commit()
    return None


@router.post("/{kb_id}/documents", response_model=DocumentUploadResponse)
async def upload_document(
    kb_id: str,
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user_required),
    db: Session = Depends(get_db),
):
    """上传文档到知识库"""
    try:
        kb_uuid = UUID(kb_id)
    except ValueError:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="无效的知识库ID格式")

    # 验证知识库存在
    kb = (
        db.query(KnowledgeBase)
        .filter(KnowledgeBase.id == kb_uuid, KnowledgeBase.user_id == current_user.id)
        .first()
    )

    if not kb:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="知识库不存在")

    # 验证文件类型
    ext = get_file_extension(file.filename)
    if ext not in ALLOWED_EXTENSIONS:
        # C72: sorted() for deterministic error messages (same as attachment_router)
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"不支持的文件类型: {ext}，支持的类型: {', '.join(sorted(ALLOWED_EXTENSIONS))}",
        )

    # 保存文件到临时目录
    # C4: strip directory components (incl. ../) from client-supplied name before
    # joining to UPLOAD_DIR — prevents path traversal writes/reads/deletes.
    safe_name = os.path.basename(file.filename or "upload")
    file_path = os.path.join(UPLOAD_DIR, f"{kb_uuid}_{safe_name}")
    # C4: defense-in-depth — reject if the resolved path escapes UPLOAD_DIR
    if not os.path.abspath(file_path).startswith(os.path.abspath(UPLOAD_DIR) + os.sep):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="无效的文件名")
    try:
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"文件保存失败: {str(e)}"
        )

    # 获取文件大小
    file_size = os.path.getsize(file_path)

    # 创建文档记录
    doc = Document(
        knowledge_base_id=kb_uuid,
        user_id=current_user.id,
        filename=file.filename,
        file_type=ext[1:] if ext else None,  # 去掉点
        file_size=file_size,
        file_path=file_path,
        status="pending",
    )
    db.add(doc)

    # 更新知识库文档计数
    kb.document_count = (kb.document_count or 0) + 1

    db.commit()
    db.refresh(doc)

    # 获取数据库会话工厂
    from app.core.database import SessionLocal

    # 在后台处理文档
    # C73: pass str(kb.id) (UUID) instead of kb.name to avoid cross-user collection collisions
    background_tasks.add_task(process_document, str(doc.id), file_path, str(kb.id), SessionLocal)

    return DocumentUploadResponse(
        id=str(doc.id),
        filename=doc.filename,
        process_status="pending",
        message="文档已上传，正在后台处理中",
    )


@router.get("/{kb_id}/documents", response_model=list[DocumentResponse])
async def get_documents(
    kb_id: str,
    current_user: User = Depends(get_current_user_required),
    db: Session = Depends(get_db),
):
    """获取知识库的文档列表"""
    try:
        kb_uuid = UUID(kb_id)
    except ValueError:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="无效的知识库ID格式")

    # 验证知识库存在
    kb = (
        db.query(KnowledgeBase)
        .filter(KnowledgeBase.id == kb_uuid, KnowledgeBase.user_id == current_user.id)
        .first()
    )

    if not kb:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="知识库不存在")

    documents = (
        db.query(Document)
        .filter(Document.knowledge_base_id == kb_uuid)
        .order_by(Document.created_at.desc())
        .all()
    )

    return [doc_to_response(doc) for doc in documents]


@router.get("/{kb_id}/documents/{doc_id}/chunks")
async def get_document_chunks(
    kb_id: str,
    doc_id: str,
    current_user: User = Depends(get_current_user_required),
    db: Session = Depends(get_db),
):
    """获取文档的所有切片"""
    from app.service.milvus_service import get_milvus_service

    try:
        kb_uuid = UUID(kb_id)
        doc_uuid = UUID(doc_id)
    except ValueError:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="无效的ID格式")

    # 验证知识库存在
    kb = (
        db.query(KnowledgeBase)
        .filter(KnowledgeBase.id == kb_uuid, KnowledgeBase.user_id == current_user.id)
        .first()
    )

    if not kb:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="知识库不存在")

    # 获取文档
    doc = (
        db.query(Document)
        .filter(Document.id == doc_uuid, Document.knowledge_base_id == kb_uuid)
        .first()
    )

    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="文档不存在")

    if doc.status != "completed":
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="文档尚未处理完成")

    # 从 Milvus 获取切片
    # C73: use KB UUID to derive the collection name (matches process_document)
    collection_name = f"kb_{kb.id}".replace("-", "_")
    print(f"[get_document_chunks] 查询切片: collection={collection_name}, filename={doc.filename}")

    try:
        milvus = get_milvus_service()
        chunks = milvus.get_chunks_by_filename(collection_name, doc.filename)
        print(f"[get_document_chunks] 找到 {len(chunks)} 个切片")
    except Exception as e:
        print(f"[get_document_chunks] Milvus 查询失败: {e}")
        # 返回空结果而不是报错
        chunks = []

    return {
        "document_id": str(doc.id),
        "filename": doc.filename,
        "chunk_count": len(chunks),
        "chunks": [
            {
                "index": chunk.get("chunk_index", i),
                "content": chunk.get("content", ""),
            }
            for i, chunk in enumerate(chunks)
        ],
    }


@router.delete("/{kb_id}/documents/{doc_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_document(
    kb_id: str,
    doc_id: str,
    current_user: User = Depends(get_current_user_required),
    db: Session = Depends(get_db),
):
    """删除文档"""
    try:
        kb_uuid = UUID(kb_id)
        doc_uuid = UUID(doc_id)
    except ValueError:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="无效的ID格式")

    # 验证知识库存在
    kb = (
        db.query(KnowledgeBase)
        .filter(KnowledgeBase.id == kb_uuid, KnowledgeBase.user_id == current_user.id)
        .first()
    )

    if not kb:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="知识库不存在")

    # 获取文档
    doc = (
        db.query(Document)
        .filter(Document.id == doc_uuid, Document.knowledge_base_id == kb_uuid)
        .first()
    )

    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="文档不存在")

    # 删除文件（如果存在）
    if doc.file_path and os.path.exists(doc.file_path):
        os.remove(doc.file_path)

    # 更新知识库文档计数
    kb.document_count = max((kb.document_count or 0) - 1, 0)

    db.delete(doc)
    db.commit()
    return None
